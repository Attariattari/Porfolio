from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path(r"E:\Porfolio\docs\Professional-AI-Topic-System-Specification.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

NAVY = "17365D"; BLUE = "2E74B5"; DARK = "1F4D78"; INK = "1F2937"
MUTED = "64748B"; PALE = "E8EEF5"; LIGHT = "F4F6F9"; GOLD = "B07D18"; WHITE = "FFFFFF"

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5); sec.page_height = Inches(11)
sec.top_margin = sec.bottom_margin = Inches(0.82)
sec.left_margin = sec.right_margin = Inches(1.0)
sec.header_distance = sec.footer_distance = Inches(0.42)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"; normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri"); normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(10.5); normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.18

for name, size, color, before, after in [
    ("Title", 30, NAVY, 0, 8), ("Subtitle", 14, MUTED, 0, 14),
    ("Heading 1", 17, BLUE, 18, 9), ("Heading 2", 13.5, BLUE, 14, 7), ("Heading 3", 11.5, DARK, 10, 5),
]:
    s = styles[name]
    s.font.name = "Calibri"; s._element.rPr.rFonts.set(qn("w:ascii"), "Calibri"); s._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    s.font.size = Pt(size); s.font.color.rgb = RGBColor.from_string(color)
    s.font.bold = name not in ["Subtitle"]
    s.paragraph_format.space_before = Pt(before); s.paragraph_format.space_after = Pt(after)
    s.paragraph_format.keep_with_next = True

for name in ["List Bullet", "List Number"]:
    s = styles[name]; s.font.name = "Calibri"; s.font.size = Pt(10.5)
    s.paragraph_format.left_indent = Inches(0.375); s.paragraph_format.first_line_indent = Inches(-0.188)
    s.paragraph_format.space_after = Pt(4); s.paragraph_format.line_spacing = 1.18

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = tcPr.find(qn("w:shd"))
    if shd is None: shd = OxmlElement("w:shd"); tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None: tcMar = OxmlElement("w:tcMar"); tcPr.append(tcMar)
    for m, v in (("top",top),("start",start),("bottom",bottom),("end",end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None: node = OxmlElement(f"w:{m}"); tcMar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")

def set_table_geometry(table, widths):
    table.autofit = False; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW")); tblW.set(qn("w:w"), str(sum(widths))); tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None: tblInd = OxmlElement("w:tblInd"); tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "120"); tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(w)); grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths[i]/1440); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER; cell_margins(cell)
            tcW = cell._tc.get_or_add_tcPr().find(qn("w:tcW")); tcW.set(qn("w:w"), str(widths[i])); tcW.set(qn("w:type"), "dxa")

def table(headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
    set_table_geometry(t, widths)
    for i, h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h; shade(c, PALE)
        for r in c.paragraphs[0].runs: r.bold=True; r.font.color.rgb=RGBColor.from_string(NAVY); r.font.size=Pt(9.5)
    for row in rows:
        cells=t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text=str(val)
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after=Pt(1); p.paragraph_format.line_spacing=1.05
                for r in p.runs: r.font.size=Pt(9.2)
    set_table_geometry(t, widths)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)
    return t

def h1(text): doc.add_heading(text, 1)
def h2(text): doc.add_heading(text, 2)
def h3(text): doc.add_heading(text, 3)
def p(text, bold_prefix=None):
    para=doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        para.add_run(bold_prefix).bold=True; para.add_run(text[len(bold_prefix):])
    else: para.add_run(text)
    return para
def bullets(items):
    for x in items: doc.add_paragraph(x, style="List Bullet")
def numbers(items):
    for x in items: doc.add_paragraph(x, style="List Number")
def callout(label, text, fill=LIGHT):
    t=doc.add_table(rows=1, cols=1); t.style="Table Grid"; set_table_geometry(t,[9360]); shade(t.cell(0,0),fill)
    q=t.cell(0,0).paragraphs[0]; q.paragraph_format.space_after=Pt(0)
    q.add_run(label+": ").bold=True; q.add_run(text)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)

# Header/footer
header=sec.header.paragraphs[0]; header.text="MUHYO TECH  |  PROFESSIONAL AI EDITORIAL SYSTEM"; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT
for r in header.runs: r.font.size=Pt(8); r.font.bold=True; r.font.color.rgb=RGBColor.from_string(MUTED)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
run=footer.add_run("Professional AI Topic System  |  Page "); run.font.size=Pt(8); run.font.color.rgb=RGBColor.from_string(MUTED)
fld=OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE"); footer._p.append(fld)

# Cover
doc.add_paragraph().paragraph_format.space_after=Pt(80)
k=doc.add_paragraph(); k.alignment=WD_ALIGN_PARAGRAPH.CENTER; rr=k.add_run("SYSTEM SPECIFICATION"); rr.bold=True; rr.font.size=Pt(10); rr.font.color.rgb=RGBColor.from_string(GOLD)
title=doc.add_paragraph(style="Title"); title.alignment=WD_ALIGN_PARAGRAPH.CENTER; title.add_run("Professional AI Topic System")
sub=doc.add_paragraph(style="Subtitle"); sub.alignment=WD_ALIGN_PARAGRAPH.CENTER; sub.add_run("Core clusters, standalone authority content, verified trends, safe preemption and evidence-controlled publishing")
doc.add_paragraph().paragraph_format.space_after=Pt(44)
callout("Purpose", "Existing Pillar-first topic flow ko preserve karte hue professional topic coverage expand karna, repetition control karna aur sirf authentic official trends ko evidence-based articles mein convert karna.", "EEF4FA")
doc.add_paragraph().paragraph_format.space_after=Pt(24)
meta=doc.add_paragraph(); meta.alignment=WD_ALIGN_PARAGRAPH.CENTER
mr=meta.add_run("Prepared for Muhyo Tech\nEditorial Automation & Content Authority Architecture\nAugust 2026"); mr.font.size=Pt(10); mr.font.color.rgb=RGBColor.from_string(MUTED)
doc.add_page_break()

h1("Executive Summary")
p("New system current topic generator ko replace nahi karega. Existing Core Web Engineering ka Pillar-first workflow bilkul preserve rahega. Iske upar professional standalone categories, verified trend discovery, three verification gates, rolling distribution, repetition protection aur safe interruption/resume capability add hogi.")
callout("Three engines", "Core Cluster Engine existing 1 Pillar + 2 Supporting flow handle karega; Standalone Authority Engine remaining professional categories ke single detailed articles banayega; Verified Trend Engine official evidence ke saath time-sensitive articles handle karega.")

h1("1. Content Distribution")
table(["Category","Rolling target","Article structure"],[
    ("Core Web Engineering","30%","1 Pillar + 2 Supporting"),("Software Architecture","15%","1 standalone authority blog"),
    ("SaaS & Product Engineering","15%","1 standalone authority blog"),("Cloud, DevOps & Reliability","10%","1 standalone authority blog"),
    ("AI for Software Development","10%","1 standalone authority blog"),("Technical SEO & Web Growth","8%","1 standalone authority blog"),
    ("UI/UX & Accessibility","7%","1 standalone authority blog"),("Verified Technology Trends","5%","1 verified standalone authority blog")
],[3600,1700,4060])
p("Percentages strict daily quota nahi hongi. Yeh rolling content targets hain. System weak topic sirf target complete karne ke liye accept nahi karega; quality gates distribution se zyada important rahenge.")

h1("2. Core Cluster Engine — 30%")
p("Sirf Core Web Engineering mein existing sequence same rahega: Pillar → Supporting 1 → Supporting 2 → next Core cluster.")
h2("Covered areas")
bullets(["Next.js, React, Node.js, MERN and MongoDB", "APIs, authentication and full-stack architecture", "Web performance, web security and deployment fundamentals"])
h2("Example cluster")
table(["Stage","Example"],[
    ("Pillar","Complete Guide to Production Node.js API Architecture"),
    ("Supporting 1","Designing Reliable Error Contracts for Node.js APIs"),
    ("Supporting 2","Production Rate Limiting for Node.js Applications")],[1800,7560])
h2("Pillar article rules")
bullets(["Approximately 2,000–3,500 useful words, depending on genuine topic depth.","Foundations, architecture decisions, implementation approach, trade-offs and limitations.","Practical examples, common mistakes and an actionable decision framework.","Useful checklist, comparison table or FAQ only where it improves understanding.","No filler, repeated explanations or unsupported performance/business claims."])
h2("Supporting article rules")
bullets(["Approximately 900–1,200 useful words around one narrow practical problem.","Parent Pillar ko repeat nahi karega; it will extend one specific area.","Parent Pillar must exist as a real database blog before a Supporting topic is selected.","Supporting 2 will not be selected before Supporting 1 is completed.","Parent and child articles will use verified internal links."])
h2("Core protection")
bullets(["Next Pillar cannot bypass an incomplete Core cluster.","Failed Pillar keeps its children blocked; rejected Pillar rejects eligible children.","Interrupted processing can recover without losing cluster position.","Used-topic history and existing duplicate protection remain preserved."])

h1("3. Standalone Authority Engine — 65%")
p("Core ke ilawa har professional category mein one topic → one complete standalone authority article hoga. In topics ke Supporting articles nahi banenge.")
callout("Required article type", "Database aur scheduler mein in articles ka type standalone_authority hona chahiye. Inko pillar kehne se existing selector unnecessary Supporting topics expect kar sakta hai.")
sections=[
    ("Software Architecture — 15%",["System design, modular architecture and modular monoliths","Microservices, event-driven systems and background jobs","Caching, multi-tenancy, data consistency, scalability and technical debt"],"Modular Monolith vs Microservices for a Growing SaaS Product"),
    ("SaaS & Product Engineering — 15%",["MVP architecture and founder technology planning","Admin dashboards, subscriptions, roles and product workflows","Build-vs-buy decisions, product scaling and feature planning"],"What Founders Should Validate Before Building an MVP"),
    ("Cloud, DevOps & Reliability — 10%",["AWS, Vercel, CI/CD and deployment previews","Monitoring, logging, observability and rollbacks","Infrastructure cost, incident recovery and environment management"],"AWS vs Vercel: Infrastructure Decisions for a Growing Product"),
    ("AI for Software Development — 10%",["LLM integrations, RAG and structured outputs","Human approval, prompt versioning and AI security","Model fallbacks, reliability and cost management"],"Why Production AI Workflows Need Deterministic Validation"),
    ("Technical SEO & Web Growth — 8%",["Crawlability, indexation, canonicals, schema and sitemaps","Core Web Vitals, JavaScript SEO and website migrations","Programmatic SEO and engineering-led growth"],"Technical SEO Checks Before Migrating a JavaScript Website"),
    ("UI/UX & Accessibility — 7%",["Dashboard UX, forms, navigation and data-heavy interfaces","Loading/error states, mobile usability and user trust","Accessibility and maintainable design systems"],"Designing Forms That Recover Gracefully from Validation Errors"),
]
for title, items, ex in sections:
    h2(title); bullets(items); p("Example: “"+ex+"”")
h2("Standalone article quality")
bullets(["Approximately 1,800–3,000 words where the topic genuinely needs that depth.","Clear audience, real problem, root cause, options and decision criteria.","Implementation considerations, trade-offs, limitations, risks and practical checklist.","Business implications stated responsibly without invented results or guarantees.","No artificial word-count padding; article ends when the problem is completely covered."])

h1("4. Professional Topic Discovery and Classification")
h2("Discovery inputs")
bullets(["Gemini-generated evergreen ideas and audience problems", "Existing content coverage gaps and under-covered categories", "Search-intent opportunities and relevant services", "Manual admin topics", "Official product blogs, changelogs, GitHub releases, advisories and platform announcements"])
h2("Classification")
p("AI assigns one category: core_web_engineering, software_architecture, saas_product_engineering, cloud_devops_reliability, ai_software_development, technical_seo_growth, uiux_accessibility or verified_trend.")
bullets(["Classification reason", "Target audience", "Real problem", "Professional value", "Muhyo Tech expertise connection", "Recommended article type"])

h1("5. General Topic Scoring")
table(["Signal","Maximum score"],[
    ("Professional audience relevance",25),("Practical usefulness",20),("Muhyo Tech expertise fit",20),("Search opportunity",15),
    ("Original angle",10),("Category coverage requirement",5),("Relevant service connection",5),("Total",100)
],[7000,2360])
callout("Acceptance threshold", "General professional topic must score at least 70/100 and pass every mandatory rule.")
h2("Mandatory rejection conditions")
bullets(["Professional audience or real problem is unclear.","Practical solution angle or credible expertise connection is missing.","Topic duplicates existing or queued content.","Topic is clickbait, temporary hype or depends on unsupported claims.","Search intent is unclear or the topic adds no meaningful value."])

h1("6. Duplicate, Repetition and Rotation Protection")
h2("Duplicate checks")
bullets(["Existing titles, summaries, focus keywords and search intent", "Existing queue, used history and rejected history", "Topic family, problem statement, solution angle and conclusion", "Target audience, category, primary technology and article format"])
p("Normalized fingerprint will combine article type, category, topic family, problem, solution angle, focus keyword and audience. Database uniqueness and semantic near-duplicate checks will both apply.")
h2("Rotation rules")
bullets(["Same primary technology will not run back-to-back.","Same topic family appears at most twice in the last five posts.","Same service focus should not repeat in the last four posts.","Same audience and article format should not repeat continuously.","Over-covered category receives a temporary cooldown; under-covered category receives a priority boost.","Similar title patterns are rejected even when exact wording differs."])
h2("Allowed article formats")
bullets(["Authority guide", "Decision framework", "Comparison", "Architecture breakdown", "Engineering deep dive", "Migration guide", "Production checklist", "Common mistakes", "Security advisory", "Performance diagnosis", "Founder’s guide", "Implementation strategy", "Case-based analysis", "Verified release impact analysis"])

h1("7. Verified Trend Engine — 5%")
p("Trend Engine random technology news generator nahi hoga. Yeh sirf professionally relevant, time-sensitive aur official evidence se verified updates cover karega.")
h2("Eligible trend families")
bullets(["Stable framework releases and Node.js LTS changes", "Official React, Next.js, MongoDB and browser-platform changes", "Security advisories and important cloud-platform changes", "Official search-platform, web-standard and AI developer-tool updates"])
h2("Always reject")
bullets(["Rumours, leaks and unverified social posts", "Gadget/mobile reviews and celebrity technology news", "Crypto speculation and generic AI hype", "Third-party-only claims without a primary official source", "Minor updates with no clear professional impact"])

h1("8. Trend Gate 1 — Authenticity Before Queue")
p("Lifecycle: Discovered → Verification Pending → Source Verified → Impact Scored → Topic Approved. Trend direct editorial queue mein enter nahi karega.")
h2("Mandatory evidence")
bullets(["Primary official source and accessible official URL", "Official title, publication date, product/framework identity", "Exact version and official release notes where applicable", "Relevant official excerpts and direct topic-to-announcement match", "No official contradiction, correction or duplicate coverage", "Clear professional impact"])
table(["Authenticity signal","Score"],[
    ("Primary official source",30),("URL accessibility",10),("Date/version confirmation",15),("Supporting official documentation",15),
    ("Claim-to-source consistency",15),("Professional relevance",10),("No contradiction/correction",5),("Total",100)
],[7000,2360])
callout("Strict rule", "Minimum authenticity score is 90/100. Primary official source is mandatory; it cannot be replaced by a high calculated score.", "FFF6E3")
h2("Evidence record")
bullets(["Source name, URL, domain, type and title", "Publication date, discovery date and verification timestamp", "Framework/product, exact version and relevant excerpts", "Verified claims, unverified claims and AI verification explanation", "Source fingerprint/hash, authenticity score and expiry date"])

h1("9. Trend Priority and Safe Preemption")
table(["Priority","Examples","Scheduler behavior"],[
    ("Critical","Critical security issue; official urgent migration","Next safe writing checkpoint"),
    ("High","Major stable/LTS release; important breaking change","After current blog completes"),
    ("Normal","Minor non-breaking or preview update","Standard editorial rotation")
],[1500,4300,3560])
p("Only verified Critical and High trends receive preemption permission. Normal categories cannot interrupt one another.")
h2("Safe checkpoint rule")
numbers(["Potential trend is discovered and verified in parallel.","Already-processing blog finishes and is safely saved; no half-written article is cancelled.","Current queue/cluster position is stored.","Trend receives the next exclusive writing slot.","Trend is re-verified, written and fact-audited.","Trend is saved as Pending Review.","Previous editorial flow resumes from the exact saved position."])
h2("Core interruption example")
p("Core Pillar completes → state saves nextStage=supporting_1 → verified trend article runs → audit passes → trend saves as Pending Review → Core resumes at Supporting 1 → Supporting 2.")
h2("Multiple trends")
bullets(["Highest criticality, authenticity, time sensitivity and audience impact wins.","Other verified trends remain verified_waiting.","Normally only one trend interrupts a normal sequence; a Critical security item may override.","Trend backlog cannot permanently starve evergreen content."])

h1("10. Editorial Run State and Concurrency")
h2("State fields")
bullets(["Active topic/category/cluster ID", "Current and last-completed cluster stage", "Next required stage", "Pause reason and interrupting trend ID", "Paused timestamp and resumeAfterTrend flag", "Processing lock and last completed blog ID"])
h2("Example state")
callout("paused_for_trend", "currentCategory=core_web_engineering; clusterKey=nodejs-api-architecture; lastCompletedStage=pillar; nextStage=supporting_1; interruptingTrendId=<verified trend>; resumeAfterTrend=true.")
h2("Concurrency locks")
bullets(["writerLock", "activeTopicId", "pausedTopicId", "interruptingTrendId", "trendVerificationLock", "processingStartedAt"])
p("Trend analysis may run in parallel, but only one article writer can run at a time. This prevents duplicate blogs, wrong slugs, cluster corruption and conflicting queue states.")

h1("11. Trend Gate 2 — Re-verification Before Writing")
bullets(["Official URL is still accessible.","Source has not materially changed.","Version and publication date still match.","No official correction or revised advisory invalidates the topic.","Central claim remains supported and the trend has not expired.","Stable/preview status is correctly represented."])
p("Failure changes the trend to verification_blocked, releases any pause and resumes the previous normal flow without generating an article.")
h2("Evidence-bound writer packet")
bullets(["Approved topic and title direction", "Official sources, exact version and dates", "Verified features, breaking changes and excerpts", "Allowed conclusions and required qualifications", "Unsupported claims and prohibited assumptions"])
callout("Writer boundary", "The AI may not add a version, date, feature, breaking change, compatibility statement, security impact, benchmark, migration requirement, deprecation, statistic or quote unless the verified evidence packet supports it.", "FFF6E3")

h1("12. Trend Gate 3 — Claim-by-Claim Audit")
p("Generated trend article will not be accepted immediately. Auditor AI extracts factual claims and checks each one against official evidence.")
table(["Audit result","Required action"],[
    ("supported","Keep the claim"),("partially_supported","Rewrite or qualify"),("unsupported","Regenerate/remove; publication blocked"),
    ("contradicted","Immediate block"),("needs_qualification","Add precise scope/limitation")
],[3000,6360])
h2("Claims inspected")
bullets(["Dates and versions", "Feature names and breaking changes", "Deprecations and compatibility", "Security and performance statements", "Migration instructions, statistics and quotes"])

h1("13. Topic-to-Writer Handoff")
bullets(["Article type and category", "Title direction and topic family", "Real problem, solution angle and business value", "Audience, focus keyword and search intent", "Recommended format and related service slugs", "Parent Pillar context where applicable", "Official evidence and prohibited claims for trends"])
p("Writer ko sirf title nahi diya jayega; complete editorial intent and safety boundary provide hogi.")

h1("14. General Article Quality Gates")
bullets(["Professional title, unique canonical slug and natural focus keyword", "Clear search intent and meaningful structure", "No duplicate sections, fake client results, guarantees or unsupported statistics", "Only relevant service links using /services/<slug>", "Only real published blog targets for /blog/<slug>", "Correct parent Pillar link where applicable", "Complete metadata, relevant image and professional readability", "No filler; generated article must be materially different from existing content"])

h1("15. Queue Status Model")
table(["General statuses","Trend-specific statuses"],[
    ("discovered","source_discovered"),("verification_pending","source_verified"),("planned","impact_approved"),("ready","verified_waiting"),
    ("processing","preemption_requested"),("paused","reverification_pending"),("used","verification_blocked"),("rejected","fact_checking"),
    ("failed","fact_check_failed"),("expired","pending_review")
],[4680,4680])

h1("16. Failure and Recovery")
h2("Normal topic failure")
bullets(["Retry count increments with a maximum of three controlled retries.","Core Pillar returns to planned; Supporting returns to ready.","After three failures, topic becomes failed while relationships remain preserved."])
h2("Trend failures")
bullets(["Authenticity failure never enters the editorial topic queue.","Pre-generation verification failure becomes verification_blocked and resumes normal flow.","Generation failure retries within limits, then moves to failed/manual review and releases the pause.","Factual audit failure triggers targeted rewrite; persistent failure requires manual review."])

h1("17. Admin Dashboard")
h2("Information shown")
bullets(["Topic title, category, article type and cluster position", "Source, professional score, authenticity score and impact score", "Priority, audience, search intent, services and schedule", "Duplicate/rotation status and Parent Pillar", "Generated blog, trend badge, version and official sources", "Verification/expiry dates, failure reason and pause/resume state"])
h2("Admin actions")
bullets(["Approve, edit, reject, retry and reschedule", "Change priority", "Re-verify a trend and view its evidence", "Mark expired", "Generate/open a blog", "Resume a paused sequence"])

h1("18. Publishing Policy")
h2("Evergreen professional article")
p("Topic Approved → Article Generated → Quality Audit → Pending Admin Review → Published")
h2("Verified trend article")
p("Trend Discovered → Official Authenticity Verification → Impact Approval → Editorial Queue → Source Re-verification → Evidence-bound Writing → Claim Audit → Pending Admin Review → Published")
callout("Resume rule", "Normal paused flow resumes once the trend article passes audit and is safely saved as Pending Review. It does not wait for admin publication.")

h1("19. Scheduler Decision Order")
numbers(["Allow the current processing job to complete.","Select a Critical verified trend if one has requested preemption.","Otherwise select a High verified trend.","Resume a previously paused sequence when the trend slot is complete.","Continue an incomplete Core cluster in Pillar → Supporting 1 → Supporting 2 order.","Select scheduled standalone authority content.","Prefer the best-scoring topic from an under-covered category."])
h2("Checks before every writing slot")
bullets(["Active processing lock", "Verified trend priority", "Paused state", "Incomplete Core cluster", "Rolling category coverage", "Topic score and rotation rules", "Schedule eligibility and duplicate safety"])

h1("20. Complete System Architecture")
p("Topic Discovery feeds three specialist engines. Core Cluster Engine produces Pillar + two Supporting topics; Standalone Authority Engine produces one detailed authority topic; Verified Trend Engine performs official-source verification and may request safe preemption. All approved items enter one Editorial Queue, then the Priority Scheduler selects one writer job. Quality or factual audits gate Pending Admin Review and publication. Any paused sequence resumes from its saved checkpoint.")
table(["Engine","Primary responsibility","Output"],[
    ("Core Cluster Engine","Preserve topical depth and strict parent-child order","1 Pillar + 2 Supporting"),
    ("Standalone Authority Engine","Expand professional coverage without unnecessary clusters","1 detailed authority article"),
    ("Verified Trend Engine","Discover and verify time-sensitive official updates","1 evidence-controlled authority article")
],[2600,4200,2560])

h1("21. Final Recommendations")
numbers(["Keep the existing Core cluster workflow as a protected editorial lane.","Add standalone_authority as a separate article type so Core selection remains unambiguous.","Use official-source code checks as the authority; AI analyzes meaning, relevance and claims.","Allow trend preemption only at a safe checkpoint after the current blog is saved.","Treat Pending Review save as trend completion so normal content does not remain blocked.","Use category percentages as rolling targets, never as permission to accept weak topics.","Run one writer at a time while allowing verification/analysis jobs in parallel.","Preserve a complete evidence record for every trend article.","Limit preemption to verified Critical and High trends; normal trends follow standard rotation.","Never allow a trend failure to permanently block the Core or authority queues."])

h1("22. Final Expected Behaviour")
bullets(["Core topics build deep topical authority through Pillar and Supporting relationships.","Standalone categories add professional variety without producing unnecessary Supporting posts.","Rotation and similarity checks prevent repetitive technology coverage.","Only authentic, officially evidenced trends enter the queue.","Trend facts are checked before the queue, before writing and after writing.","Only verified Critical/High trends temporarily pause normal sequencing.","After a trend is safely saved, the system resumes from the exact previous topic position.","Fake releases, invented versions, wrong dates and unsupported claims are blocked before publication."])

doc.core_properties.title = "Professional AI Topic System Specification"
doc.core_properties.subject = "Editorial topic architecture, verified trends and safe preemption"
doc.core_properties.author = "Muhyo Tech"
doc.core_properties.keywords = "AI topic system, editorial automation, verified trends, pillar content"
doc.save(OUT)
print(OUT)
